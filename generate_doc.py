import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_project_documentation():
    # 1. Initialize Document
    doc = Document()

    # Title
    title = doc.add_heading('Project Documentation: Startup Funding Intelligence Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Overview
    doc.add_heading('1. Project Overview & Use Cases', level=1)
    doc.add_paragraph(
        "The Startup Funding Intelligence Platform is a full-stack, AI-powered web platform built to track, analyze, "
        "and predict startup funding dynamics, with a specific focus on the Indian Startup Ecosystem. It combines real-time data "
        "scraping, mathematical predictive modeling, and interactive data visualizations to offer comprehensive insights for investors and entrepreneurs."
    )
    doc.add_heading('Key Functionalities:', level=2)
    doc.add_paragraph("• Predictive Modeling: A custom-built machine learning model that forecasts the likelihood of a startup securing funding based on features like industry, geographical location, founding year, and team size.", style='List Bullet')
    doc.add_heading('Use Cases:', level=2)
    doc.add_paragraph("• For Venture Capitalists: Quickly filter the landscape for emerging startups and utilize algorithmic likelihood metrics as a proxy for investment diligence.", style='List Bullet')
    doc.add_paragraph("• For Founders: Benchmarking an upcoming or established startup against successful industry peers (unicorn modeling).", style='List Bullet')
    doc.add_paragraph("• For Market Analysts: Uncovering meta-trends in the global and local startup climate by interacting with dynamic distribution metrics.", style='List Bullet')

    # 2. Technology Stack
    doc.add_heading('2. Technology Stack', level=1)
    
    doc.add_heading('Frontend:', level=2)
    doc.add_paragraph("• React (Vite): Provides a lightning-fast modern user interface.", style='List Bullet')
    doc.add_paragraph("• Recharts: Renders responsive, dynamic SVGs mapping the underlying analytical JSON payloads.", style='List Bullet')
    doc.add_paragraph("• Vanilla CSS: Adopts a premium, glassmorphism UI/UX aesthetic prioritizing deep contrast and micro-animations.", style='List Bullet')

    doc.add_heading('Backend & API:', level=2)
    doc.add_paragraph("• FastAPI (Python): Powers ultra-fast, asynchronous endpoints for UI data consumption and live inference routing.", style='List Bullet')
    doc.add_paragraph("• Uvicorn: ASGI web server fielding concurrent frontend requests.", style='List Bullet')

    doc.add_heading('Machine Learning & Data Processing:', level=2)
    doc.add_paragraph("• XGBoost: Advanced gradient-boosting tree classifier executing the core predictor logic.", style='List Bullet')
    doc.add_paragraph("• Scikit-Learn: Employs target encoding, scaling, and pipeline vectorization on categorical attributes.", style='List Bullet')
    doc.add_paragraph("• MLflow: Records model experiments, traces hyperparameters, and snapshots metric history over training iterations.", style='List Bullet')
    doc.add_paragraph("• BeautifulSoup4 & Requests: Constructs web spiders for scraping real-time media feeds.", style='List Bullet')

    doc.add_heading('Infrastructure:', level=2)
    doc.add_paragraph("• Docker Compose: Containerizes all application dependencies into identical isolated environments.", style='List Bullet')
    doc.add_paragraph("• Supabase (Architecture provisioned): Ready-state configuration for a managed PostgreSQL cloud deployment.", style='List Bullet')

    # 3. Data Pipeline & Working Mechanism
    doc.add_heading('3. Data Flow & Working Mechanism', level=1)
    
    doc.add_heading('Where the data is fetched:', level=2)
    doc.add_paragraph("1. Historical Ecosystem Generation: A specialized `seed_data.py` algorithm mathematically reproduces 1000+ realistic permutations of Indian startups, utilizing actual market matrices (e.g. higher success probability in hubs like Bengaluru or Mumbai). It includes high-profile data signatures like Flipkart and Paytm.")
    doc.add_paragraph("2. Real-time RSS Feed Scraping: The platform contains live web agents pointing at major news hubs like TechCrunch and Inc42. The system executes complex RegEx protocols to instantly extract the target startup name, precise funding figures (millions/billions), and funding round types immediately as they break via headline.")

    doc.add_heading('Where it is stored:', level=2)
    doc.add_paragraph("In localized configuration, the data is consolidated into an integrated CSV (`training_data.csv`). For deployed production contexts, the structured payload interfaces via a python SDK to a cloud-based Supabase PostgreSQL database.")

    doc.add_heading('How it is accessed and displayed (The Connections):', level=2)
    doc.add_paragraph("1. Model Connection: The backend train_model.py script routinely loads the data warehouse. It partitions it, trains the internal XGBoost system, exports algorithmic binaries (.pkl files), and relays metric data sequentially to the MLFlow Tracking server.")
    doc.add_paragraph("2. API Connection: The FastAPI server instantiates the finalized model instance into server memory. Simultaneously, it constructs API routes (`/startups`, `/industry-trends`) which execute aggregated dataframe manipulations to synthesize the millions of startup datapoints into consumable JSON formats.")
    doc.add_paragraph("3. Client Connection: The React UI continuously queries these backend paths asynchronously. If a user utilizes the Prediction Dashboard, their inputs are serialized and dispatched locally to the FastAPI routing which calculates raw probability via the XGBoost artifact and returns color-calibrated percentage likelihoods overlaying the UI.")

    # 4. Project Quality & Potential
    doc.add_heading('4. Quality & Future Enhancements', level=1)

    doc.add_heading('Quality Indicators:', level=2)
    doc.add_paragraph("• Separation of Concerns: Strict detachment of raw ML architecture, networking (FastAPI endpoints), UI mapping, and web-scraping components into modular codebases.", style='List Bullet')
    doc.add_paragraph("• Containerization Standards: Highly scalable microservice design enforced natively by pre-configured Docker definitions.", style='List Bullet')
    doc.add_paragraph("• Defensive Data Engineering: The backend prevents API instability utilizing Fallback/Retry methodologies and integrated Python error logging.", style='List Bullet')

    doc.add_heading('Future Enhancements:', level=2)
    doc.add_paragraph("1. Automated Retraining: Architect a continuous-integration pipeline that triggers model rebases whenever the web scraper harvests >500 new real-time funding events.")
    doc.add_paragraph("2. Advanced NLP Processing: Integrate a transformer-based Large Language Model (e.g., Llama/GPT) to interpret raw textual descriptions of the startups and embed that context straight into the XGBClassifier via vector embeddings.")
    doc.add_paragraph("3. User Accounts & Portfolios: Engage the pre-wired Supabase Authentication protocols to allow VCs to 'bookmark' startups and run private 'what-if' analyses on particular founders.")

    # Save Document
    file_path = "Startup_Funding_Intelligence_Documentation.docx"
    doc.save(file_path)
    print(f"Documentation generated successfully at: {file_path}")

if __name__ == "__main__":
    create_project_documentation()
